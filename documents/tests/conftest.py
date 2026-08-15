"""Fixtures + minimal binary-document builders for the documents plugin tests."""

from __future__ import annotations

from pathlib import Path

import pytest


def make_simple_pdf(text: str = "Hello Xyberos PDF") -> bytes:
    """Build a minimal single-page PDF with one text line (valid xref table)."""
    stream = b"BT /F1 20 Tf 72 720 Td (" + text.encode("latin-1", errors="replace") + b") Tj ET"
    objects = {
        1: b"<< /Type /Catalog /Pages 2 0 R >>",
        2: b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        3: (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>"
        ),
        4: b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
        5: b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets: dict[int, int] = {}
    for number in sorted(objects):
        offsets[number] = len(out)
        out += f"{number} 0 obj\n".encode()
        out += objects[number] + b"\nendobj\n"
    xref_position = len(out)
    size = len(objects) + 1
    out += f"xref\n0 {size}\n".encode()
    out += b"0000000000 65535 f \n"
    for number in sorted(objects):
        out += f"{offsets[number]:010d} 00000 n \n".encode()
    out += f"trailer\n<< /Size {size} /Root 1 0 R >>\nstartxref\n{xref_position}\n%%EOF\n".encode()
    return bytes(out)


@pytest.fixture()
def sample_txt(tmp_path: Path) -> Path:
    path = tmp_path / "notes.md"
    path.write_text("The quick brown fox jumps over the lazy dog.", encoding="utf-8")
    return path


@pytest.fixture()
def sample_dir(tmp_path: Path) -> Path:
    """A small directory tree mixing text and binary document types."""
    (tmp_path / "a.md").write_text("Alpha document text.", encoding="utf-8")
    (tmp_path / "b.txt").write_text("Bravo plain text.", encoding="utf-8")
    (tmp_path / "page.html").write_text(
        "<html><head><title>Page</title></head><body><p>Hello HTML</p></body></html>",
        encoding="utf-8",
    )
    nested = tmp_path / "nested"
    nested.mkdir()
    (nested / "c.md").write_text("Charlie nested text.", encoding="utf-8")
    hidden = nested / ".hidden"
    hidden.mkdir()
    (hidden / "skip.md").write_text("should not load", encoding="utf-8")
    return tmp_path


@pytest.fixture()
def sample_pdf(tmp_path: Path) -> Path:
    path = tmp_path / "sample.pdf"
    path.write_bytes(make_simple_pdf("Hello Xyberos PDF"))
    return path


@pytest.fixture()
def sample_docx(tmp_path: Path) -> Path:
    pytest.importorskip("docx")
    from docx import Document

    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "1"
    doc.save(str(path))
    return path


@pytest.fixture()
def sample_xlsx(tmp_path: Path) -> Path:
    pytest.importorskip("openpyxl")
    from openpyxl import Workbook

    path = tmp_path / "sample.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Sheet1"
    sheet.append(["name", "value"])
    sheet.append(["alpha", 1])
    sheet.append(["beta", 2])
    workbook.save(str(path))
    return path

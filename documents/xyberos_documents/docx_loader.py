"""DOCX loader — lazy ``python-docx`` import.

Extracts paragraphs and tables. Raises a clear
:class:`~xyberos.exceptions.provider.ProviderError` when ``python-docx`` is not
installed (``pip install xyberos[documents]``).
"""

from __future__ import annotations

from pathlib import Path

from xyberos.exceptions.provider import ProviderError

from .base import Document, chunk_documents


def _read_docx(path: Path) -> tuple[str, int]:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise ProviderError(
            "the 'python-docx' package is required to load DOCX files; "
            "install it with 'pip install xyberos[documents]'"
        ) from exc

    doc = DocxDocument(str(path))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    table_count = 0
    for table in doc.tables:
        table_count += 1
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts), table_count


class DocxLoader:
    """Loads a Word ``.docx`` file as text (paragraphs + tables)."""

    def __init__(self, chunk_size: int | None = None) -> None:
        self._chunk_size = chunk_size

    def load(self, path: str) -> list[Document]:
        path_obj = Path(path)
        if not path_obj.is_file():
            raise FileNotFoundError(path)
        text, table_count = _read_docx(path_obj)
        document = Document(
            source=str(path_obj),
            text=text,
            metadata={"tables": table_count, "extension": ".docx"},
        )
        return chunk_documents([document], self._chunk_size)

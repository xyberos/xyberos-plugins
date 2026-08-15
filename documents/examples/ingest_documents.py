"""Example (M1): ingest a real PDF + DOCX into an IngestingKnowledge.

Builds a tiny knowledge app backed by SQLite + a deterministic embedder, loads
the documents plugin, generates a sample PDF and DOCX on the fly, and ingests
them. Run from this folder:

    python examples/ingest_documents.py

Requires the ``[documents]`` extras for PDF/DOCX (``pypdf`` / ``python-docx``).
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

PLUGIN_ROOT = Path(__file__).resolve().parents[1]
if str(PLUGIN_ROOT) not in sys.path:
    sys.path.insert(0, str(PLUGIN_ROOT))

from xyberos import create_app
from xyberos.knowledge import IngestingKnowledge
from xyberos.llm import HashEmbedder
from xyberos.vector import SqliteVectorStore

from xyberos_documents import DocumentsPlugin, load_document


def make_sample_pdf(path: Path) -> None:
    """Write a minimal single-page PDF containing one text line."""
    text = "Xyberos PDF loader reads this line."
    stream = b"BT /F1 20 Tf 72 720 Td (" + text.encode("latin-1", errors="replace") + b") Tj ET"
    objects = {
        1: b"<< /Type /Catalog /Pages 2 0 R >>",
        2: b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        3: (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>"
        ),
        4: b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
        5: b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    }
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for number in sorted(objects):
        offsets[number] = len(out)
        out += f"{number} 0 obj\n".encode() + objects[number] + b"\nendobj\n"
    xref_position = len(out)
    size = len(objects) + 1
    out += f"xref\n0 {size}\n".encode() + b"0000000000 65535 f \n"
    for number in sorted(objects):
        out += f"{offsets[number]:010d} 00000 n \n".encode()
    out += f"trailer\n<< /Size {size} /Root 1 0 R >>\nstartxref\n{xref_position}\n%%EOF\n".encode()
    path.write_bytes(bytes(out))


def make_sample_docx(path: Path) -> None:
    """Write a small DOCX (requires python-docx)."""
    from docx import Document

    doc = Document()
    doc.add_heading("Xyberos DOCX loader", level=1)
    doc.add_paragraph("The Xyberos DOCX loader reads paragraphs and tables.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Component"
    table.cell(0, 1).text = "Status"
    table.cell(1, 0).text = "DOCX"
    table.cell(1, 1).text = "works"
    doc.save(str(path))


def main() -> None:
    knowledge = IngestingKnowledge(SqliteVectorStore(":memory:"), embedder=HashEmbedder())
    app = create_app(knowledge=knowledge)
    app.load_plugin(DocumentsPlugin())

    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        sample_pdf = tmp_dir / "sample.pdf"
        sample_docx = tmp_dir / "sample.docx"
        make_sample_pdf(sample_pdf)
        make_sample_docx(sample_docx)

        for sample in (sample_pdf, sample_docx):
            docs = load_document(str(sample))
            print(f"\n{sample.name}: {len(docs)} document(s) loaded")
            for doc in docs:
                print(f"  metadata={doc.metadata}")
                print(f"  preview={doc.text[:80]!r}")
            result = app.tools.execute("ingest_document", None, path=str(sample))
            print(f"  ingested -> {result}")

    print("\nknowledge query for a known phrase:")
    print(app.knowledge.query(type("Ctx", (), {"prompt": "reads paragraphs and tables"})()))

    app.unload_plugin("documents")


if __name__ == "__main__":
    main()
